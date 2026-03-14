from docx import Document
from docx.shared import Pt
import io

def generate_docx(resume_json: dict) -> io.BytesIO:
    doc = Document()
    
    # Stylistic choices
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Contact Info
    contact = resume_json.get("contact", {})
    if contact.get("name"):
        name_p = doc.add_heading(contact["name"], 0)
        name_p.alignment = 1 # Center

    contact_details = []
    if contact.get("email"): contact_details.append(contact["email"])
    if contact.get("phone"): contact_details.append(contact["phone"])
    if contact.get("location"): contact_details.append(contact["location"])

    if contact_details:
        p = doc.add_paragraph(" | ".join(contact_details))
        p.alignment = 1

    if contact.get("linkedin"):
        p = doc.add_paragraph(contact["linkedin"])
        p.alignment = 1

    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(resume_json.get("summary", ""))

    # Skills
    skills = resume_json.get("skills", [])
    if isinstance(skills, dict):
        derived = []
        derived.extend(skills.get("explicit") or [])
        derived.extend(skills.get("inferred_from_experience_projects") or [])
        grouped = skills.get("grouped") or {}
        if isinstance(grouped, dict):
            for items in grouped.values():
                if items:
                    derived.extend(items)
        # Deduplicate while preserving order
        seen = set()
        skills = [s for s in derived if isinstance(s, str) and not (s in seen or seen.add(s))]
    elif isinstance(skills, str):
        skills = [s.strip() for s in skills.split(",") if s.strip()]
    if skills:
        doc.add_heading('Key Skills', level=1)
        doc.add_paragraph(", ".join(skills))

    # Experience
    doc.add_heading('Work Experience', level=1)
    for exp in resume_json.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('title', 'N/A')} - {exp.get('company', 'N/A')}")
        run.bold = True
        p_period = doc.add_paragraph()
        run_period = p_period.add_run(exp.get("period", ""))
        run_period.italic = True
        for bullet in exp.get("bullets", []):
            doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('Education', level=1)
    for edu in resume_json.get("education", []):
        p = doc.add_paragraph()
        degree_text = edu.get('degree', 'N/A')
        field = edu.get('field_of_study', '')
        if field:
            degree_text = f"{degree_text} — {field}"
        run = p.add_run(degree_text)
        run.bold = True
        doc.add_paragraph(f"{edu.get('school', 'N/A')} ({edu.get('year', 'N/A')})")

    # Certifications
    certs = resume_json.get("certifications", [])
    if certs:
        doc.add_heading('Certifications', level=1)
        for cert in certs:
            p = doc.add_paragraph()
            if isinstance(cert, str):
                run = p.add_run(cert)
                run.bold = True
            else:
                run = p.add_run(cert.get('name', 'N/A'))
                run.bold = True
                issuer = cert.get('issuer', '')
                date = cert.get('date', '')
                detail_parts = [x for x in [issuer, date] if x]
                if detail_parts:
                    doc.add_paragraph(" | ".join(detail_parts))

    # Projects
    projects = resume_json.get("projects", [])
    if projects:
        doc.add_heading('Projects', level=1)
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get('name', 'N/A'))
            run.bold = True
            if proj.get('description'):
                doc.add_paragraph(proj['description'])
            tech = proj.get('tech_stack', [])
            if tech:
                doc.add_paragraph(f"Technologies: {', '.join(tech)}")
            for outcome in proj.get('outcomes', []):
                doc.add_paragraph(outcome, style='List Bullet')

    # Publications
    publications = resume_json.get("publications", [])
    if publications:
        doc.add_heading('Publications & Talks', level=1)
        for pub in publications:
            doc.add_paragraph(pub if isinstance(pub, str) else pub.get('title', str(pub)),
                              style='List Bullet')

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
